import os
import queue
from datetime import datetime
import tkinter as tk
from tkinter import simpledialog, filedialog, messagebox

from PIL import ImageGrab, ImageDraw, ImageFont
from pynput import mouse, keyboard

from docx import Document
from docx.shared import Inches

# Optional PDF export
try:
    from docx2pdf import convert
    DOCX2PDF = True
except ImportError:
    DOCX2PDF = False

# Optional active‑window title (needs pywin32)
try:
    import win32gui  # type: ignore
except ImportError:
    win32gui = None  # fallback

# Try to import customtkinter for modern GUI
try:
    import customtkinter as ctk
    CTK_AVAILABLE = True
except ImportError:
    CTK_AVAILABLE = False
    print("customtkinter not available, using standard tkinter")
    print("Install with: pip install customtkinter")


class StepRecorderApp:
    """PSR‑style recorder with click + keystroke capture, active‑window titles, cute emoji UI, and unlimited sessions."""

    def __init__(self, master):
        self.master = master
        
        if CTK_AVAILABLE:
            # Configure customtkinter
            ctk.set_appearance_mode("dark")
            ctk.set_default_color_theme("blue")
            master.title("🎬 Step Recorder Pro")
            master.geometry("400x500")
            master.resizable(False, False)
        else:
            master.title("Step Recorder 🐾")
            master.geometry("350x450")
            master.resizable(False, False)

        # ── state ───────────────────────────────────────────────────────
        self.save_path: str | None = None
        self.recording = False
        self.step = 0
        self.queue: queue.Queue = queue.Queue()
        self.outputs: list[str] = []
        self.last_click = (100, 100)

        # ── keystroke grouping ──────────────────────────────────────────
        self.current_keystrokes = []
        self.last_key_time = 0
        self.key_grouping_delay = 1000  # 1 second delay to group keystrokes

        # ── UI ──────────────────────────────────────────────────────────
        self._create_ui()

        # listeners
        self.mouse_listener: mouse.Listener | None = None
        self.kb_listener:    keyboard.Listener | None = None

        self.doc: Document | None = None

    def _create_ui(self):
        """Create the modern UI with customtkinter or fallback to tkinter"""
        if CTK_AVAILABLE:
            self._create_modern_ui()
        else:
            self._create_standard_ui()

    def _create_modern_ui(self):
        """Create modern UI with customtkinter"""
        # Main frame
        main_frame = ctk.CTkFrame(self.master)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Title
        title_label = ctk.CTkLabel(
            main_frame, 
            text="🎬 Step Recorder Pro", 
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=(0, 20))

        # File selection section
        file_frame = ctk.CTkFrame(main_frame)
        file_frame.pack(fill="x", padx=10, pady=10)

        file_label = ctk.CTkLabel(file_frame, text="📁 Output File", font=ctk.CTkFont(size=16))
        file_label.pack(pady=(10, 5))

        self.select_file_btn = ctk.CTkButton(
            file_frame, 
            text="Choose Save Location", 
            command=self.select_save_path,
            fg_color="#2E8B57",
            hover_color="#228B22"
        )
        self.select_file_btn.pack(pady=5)

        self.path_label = ctk.CTkLabel(
            file_frame, 
            text="No file selected", 
            text_color="gray",
            font=ctk.CTkFont(size=12)
        )
        self.path_label.pack(pady=(0, 10))

        # Control section
        control_frame = ctk.CTkFrame(main_frame)
        control_frame.pack(fill="x", padx=10, pady=10)

        control_label = ctk.CTkLabel(control_frame, text="🎮 Recording Controls", font=ctk.CTkFont(size=16))
        control_label.pack(pady=(10, 15))

        # Buttons frame
        btn_frame = ctk.CTkFrame(control_frame)
        btn_frame.pack(fill="x", padx=20, pady=10)

        self.start_btn = ctk.CTkButton(
            btn_frame, 
            text="▶️ Start Recording", 
            command=self.start,
            fg_color="#4CAF50",
            hover_color="#45a049",
            height=40,
            font=ctk.CTkFont(size=14, weight="bold")
        )
        self.start_btn.pack(fill="x", pady=5)

        self.note_btn = ctk.CTkButton(
            btn_frame, 
            text="📝 Add Note", 
            command=self.add_note,
            fg_color="#FF9800",
            hover_color="#F57C00",
            height=35
        )
        self.note_btn.pack(fill="x", pady=5)

        self.stop_btn = ctk.CTkButton(
            btn_frame, 
            text="⏹️ Stop Recording", 
            command=self.stop,
            fg_color="#f44336",
            hover_color="#da190b",
            height=40,
            font=ctk.CTkFont(size=14, weight="bold")
        )
        self.stop_btn.pack(fill="x", pady=5)

        # Status section
        status_frame = ctk.CTkFrame(main_frame)
        status_frame.pack(fill="x", padx=10, pady=10)

        status_label = ctk.CTkLabel(status_frame, text="📊 Status", font=ctk.CTkFont(size=16))
        status_label.pack(pady=(10, 5))

        self.status_label = ctk.CTkLabel(
            status_frame, 
            text="Ready to record", 
            text_color="lightgreen",
            font=ctk.CTkFont(size=14)
        )
        self.status_label.pack(pady=(0, 10))

        # Instructions
        info_frame = ctk.CTkFrame(main_frame)
        info_frame.pack(fill="x", padx=10, pady=10)

        info_text = """
💡 Instructions:
• Select save location first
• Click 'Start Recording' to begin
• Press ESC to stop recording
• Add notes anytime during recording
• Screenshots are taken automatically
        """
        
        info_label = ctk.CTkLabel(
            info_frame, 
            text=info_text,
            font=ctk.CTkFont(size=12),
            justify="left"
        )
        info_label.pack(pady=10)

    def _create_standard_ui(self):
        """Create standard UI with tkinter"""
        # Title
        title_label = tk.Label(
            self.master, 
            text="Step Recorder 🐾", 
            font=("Arial", 16, "bold")
        )
        title_label.pack(pady=10)

        # File selection
        tk.Button(
            self.master, 
            text="📁 Select Save File", 
            command=self.select_save_path,
            bg="#4CAF50",
            fg="white",
            font=("Arial", 10, "bold")
        ).pack(padx=10, pady=4)
        
        self.path_label = tk.Label(
            self.master, 
            text="No file selected", 
            fg="grey",
            font=("Arial", 9)
        )
        self.path_label.pack()

        # Control buttons
        tk.Button(
            self.master, 
            text="▶️ Start Recording", 
            command=self.start,
            bg="#2196F3",
            fg="white",
            font=("Arial", 10, "bold")
        ).pack(padx=10, pady=4)
        
        tk.Button(
            self.master, 
            text="📝 Add Note", 
            command=self.add_note,
            bg="#FF9800",
            fg="white"
        ).pack(padx=10, pady=4)
        
        tk.Button(
            self.master, 
            text="⏹️ Stop Recording", 
            command=self.stop,
            bg="#f44336",
            fg="white",
            font=("Arial", 10, "bold")
        ).pack(padx=10, pady=4)

        self.status_label = tk.Label(
            self.master, 
            text="Not recording",
            font=("Arial", 10)
        )
        self.status_label.pack(pady=(8, 0))

    # ── helper: active window title ────────────────────────────────────
    @staticmethod
    def get_active_window_title() -> str:
        if win32gui:
            try:
                hwnd = win32gui.GetForegroundWindow()
                title = win32gui.GetWindowText(hwnd)
                return title or "Unknown App"
            except Exception:
                return "Unknown App"
        return "Unknown App"

    # ── UI callbacks ───────────────────────────────────────────────────
    def select_save_path(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Select output .docx file",
        )
        if path:
            self.save_path = path
            filename = os.path.basename(path)
            if CTK_AVAILABLE:
                self.path_label.configure(text=filename, text_color="white")
            else:
                self.path_label.config(text=filename, fg="black")

    def start(self):
        if self.recording:
            return
        if not self.save_path:
            messagebox.showwarning("No file", "Please select where to save first.")
            return

        # reset session
        self.recording = True
        if CTK_AVAILABLE:
            self.status_label.configure(text="● Recording…", text_color="red")
        else:
            self.status_label.config(text="● Recording…", fg="red")
        self.step = 0
        self.queue = queue.Queue()
        self.outputs = []
        self.current_keystrokes = []
        self.last_key_time = 0

        self.doc = Document()
        self.doc.add_heading("Recorded Steps", level=1)

        # start listeners with error handling
        try:
            print("Starting mouse listener...")
            self.mouse_listener = mouse.Listener(on_click=self._on_click)
            self.mouse_listener.start()
            print("Mouse listener started successfully.")
            
            print("Starting keyboard listener...")
            self.kb_listener = keyboard.Listener(on_release=self._on_key_release)
            self.kb_listener.start()
            print("Keyboard listener started successfully.")
            
        except Exception as e:
            print(f"Failed to start listeners: {e}")
            messagebox.showerror("Error", f"Failed to start listeners: {e}")
            self.recording = False
            if CTK_AVAILABLE:
                self.status_label.configure(text="Failed to start", text_color="red")
            else:
                self.status_label.config(text="Failed to start", fg="red")
            return

        # poll queue
        self.master.after(100, self._process_queue)

    def stop(self):
        if not self.recording:
            return
        
        # Record any remaining keystrokes before stopping
        if self.current_keystrokes:
            self._record_keystroke_group()
        
        self.recording = False
        if CTK_AVAILABLE:
            self.status_label.configure(text="Stopped", text_color="lightgreen")
        else:
            self.status_label.config(text="Stopped", fg="black")

        # stop listeners
        if self.mouse_listener:
            self.mouse_listener.stop(); self.mouse_listener = None
        if self.kb_listener:
            self.kb_listener.stop(); self.kb_listener = None

        # save document
        if self.save_path and self.doc:
            self.doc.save(self.save_path)
            if DOCX2PDF:
                try:
                    pdf_out = os.path.splitext(self.save_path)[0] + ".pdf"
                    convert(self.save_path, pdf_out)
                    messagebox.showinfo("Saved", f"DOCX → {self.save_path}\nPDF  → {pdf_out}")
                except Exception:
                    messagebox.showinfo("Saved", f"DOCX → {self.save_path}\n(PDF conversion failed)")
            else:
                messagebox.showinfo("Saved", f"DOCX → {self.save_path}\n(Install docx2pdf for PDF)")
        else:
            messagebox.showwarning("Not saved", "No save path; nothing written.")

        # cleanup screenshots
        for f in self.outputs:
            try:
                os.remove(f)
            except FileNotFoundError:
                pass

    def add_note(self):
        if not self.recording:
            return
        note = simpledialog.askstring("Add Note", "Enter your note:")
        if note:
            self.doc.add_paragraph(f"Note: {note}")
            self.doc.add_paragraph("")

    # ── listeners ───────────────────────────────────────────────────────
    def _on_click(self, x, y, button, pressed):
        if pressed and self.recording:
            print(f"Mouse click detected: {button} at ({x}, {y})")  # Debug print
            
            # Record any pending keystrokes before the click
            if self.current_keystrokes:
                self._record_keystroke_group()
            
            self.last_click = (x, y)
            self.queue.put(("click", x, y, button))

    def _on_key_release(self, key):
        """Capture keystrokes on release so printable char is populated."""
        print(f"Key released: {key}")  # Debug print
        if not self.recording:
            print("Not recording, ignoring key")
            return
        if key == keyboard.Key.esc:  # ESC stops recording
            print("ESC pressed, stopping recording")
            self.master.after(0, self.stop)
            return False
        
        # Add to current keystroke group
        current_time = datetime.now().timestamp() * 1000  # milliseconds
        
        # If too much time has passed since last key, record the previous group
        if self.current_keystrokes and (current_time - self.last_key_time) > self.key_grouping_delay:
            self._record_keystroke_group()
        
        # Add current key to group
        self.current_keystrokes.append(key)
        self.last_key_time = current_time
        
        # Schedule recording of this group after delay
        self.master.after(self.key_grouping_delay, self._schedule_keystroke_recording)

    def _schedule_keystroke_recording(self):
        """Schedule recording of current keystroke group if no new keys have been pressed."""
        current_time = datetime.now().timestamp() * 1000
        if self.current_keystrokes and (current_time - self.last_key_time) >= self.key_grouping_delay:
            self._record_keystroke_group()

    def _record_keystroke_group(self):
        """Record the current group of keystrokes as a single step."""
        if not self.current_keystrokes:
            return
        
        try:
            self.step += 1
            ts = datetime.now().strftime("%H:%M:%S")
            title = self.get_active_window_title()
            
            # Create a combined keystroke string
            keystroke_text = self._keystrokes_to_text(self.current_keystrokes)
            
            # Take screenshot only once for the group
            img_path = self._capture_and_highlight(*self.last_click)
            
            self.doc.add_paragraph(
                f"Step {self.step:03d} at {ts} in '{title}': Typed '{keystroke_text}'")
            self.doc.add_picture(img_path, width=Inches(5))
            self.doc.add_paragraph("")
            
        except Exception as e:
            print(f"Error recording keystroke group: {e}")
            # Try to record a simplified version
            try:
                self.step += 1
                ts = datetime.now().strftime("%H:%M:%S")
                title = self.get_active_window_title()
                
                self.doc.add_paragraph(
                    f"Step {self.step:03d} at {ts} in '{title}': Keystrokes recorded (details unavailable)")
                self.doc.add_paragraph("")
            except Exception as e2:
                print(f"Failed to record even simplified keystroke group: {e2}")
        
        # Clear the group
        self.current_keystrokes = []

    def _keystrokes_to_text(self, keystrokes):
        """Convert a list of keystrokes to a readable text string."""
        text_parts = []
        for key in keystrokes:
            key_name = self._key_to_name(key)
            # Filter out control characters and problematic keys
            if self._is_safe_for_xml(key_name) and key_name not in ['Char(19)', 'Ctrl_L', 'Shift_L', 'Alt_L']:
                text_parts.append(key_name)
        
        if not text_parts:
            # Check if we have any meaningful special keys
            special_keys = []
            for key in keystrokes:
                key_name = self._key_to_name(key)
                if key_name in ['Enter', 'Space', 'Tab', 'Backspace', 'Delete', 'Escape', 'Up Arrow', 'Down Arrow', 'Left Arrow', 'Right Arrow']:
                    special_keys.append(key_name)
            
            if special_keys:
                return f"Pressed: {', '.join(special_keys)}"
            else:
                return "Special keys pressed"
        
        return "".join(text_parts)  # Join without spaces for better readability

    def _is_safe_for_xml(self, text):
        """Check if text is safe to write to XML/Word document."""
        if not text:
            return False
        
        # Check for control characters (ASCII 0-31 except tab, newline, carriage return)
        for char in text:
            if ord(char) < 32 and char not in '\t\n\r':
                return False
            # Check for other problematic characters
            if ord(char) > 127 and not char.isprintable():
                return False
        
        return True

    # key friendly name
    @staticmethod
    def _key_to_name(key) -> str:
        if hasattr(key, "char") and key.char is not None:
            # Handle printable characters
            if key.char == " ":
                return "Space"
            elif key.char.isprintable():
                return key.char
            else:
                return f"Char({ord(key.char)})"
        
        # Handle special keys
        key_str = str(key).replace("Key.", "").replace("'", "").title()
        
        # Map common special keys to readable names
        key_mapping = {
            "Enter": "Enter",
            "Space": "Space", 
            "Tab": "Tab",
            "Backspace": "Backspace",
            "Delete": "Delete",
            "Escape": "Escape",
            "Esc": "Escape",
            "Shift": "Shift",
            "Shift_L": "Shift",
            "Shift_R": "Shift",
            "Ctrl": "Ctrl",
            "Ctrl_L": "Ctrl",
            "Ctrl_R": "Ctrl",
            "Alt": "Alt",
            "Alt_L": "Alt",
            "Alt_R": "Alt",
            "Up": "Up Arrow",
            "Down": "Down Arrow", 
            "Left": "Left Arrow",
            "Right": "Right Arrow",
            "Home": "Home",
            "End": "End",
            "Page_Up": "Page Up",
            "Page_Down": "Page Down",
            "Insert": "Insert",
            "F1": "F1",
            "F2": "F2", 
            "F3": "F3",
            "F4": "F4",
            "F5": "F5",
            "F6": "F6",
            "F7": "F7",
            "F8": "F8", 
            "F9": "F9",
            "F10": "F10",
            "F11": "F11",
            "F12": "F12"
        }
        
        return key_mapping.get(key_str, key_str)

    # ── queue processor ────────────────────────────────────────────────
    def _process_queue(self):
        while not self.queue.empty():
            kind, *data = self.queue.get()
            print(f"Processing queue item: {kind} with data {data}")  # Debug print
            if kind == "click":
                x, y, button = data
                self._record_click(x, y, button)
        if self.recording:
            self.master.after(100, self._process_queue)

    # ── record helpers ─────────────────────────────────────────────────
    def _record_click(self, x, y, button):
        try:
            self.step += 1
            ts = datetime.now().strftime("%H:%M:%S")
            title = self.get_active_window_title()
            img_path = self._capture_and_highlight(x, y)
            self.doc.add_paragraph(
                f"Step {self.step:03d} at {ts} in '{title}': Click {button} at ({x},{y})")
            self.doc.add_picture(img_path, width=Inches(5))
            self.doc.add_paragraph("")
        except Exception as e:
            print(f"Error recording click: {e}")
            # Try to record a simplified version
            try:
                self.step += 1
                ts = datetime.now().strftime("%H:%M:%S")
                title = self.get_active_window_title()
                self.doc.add_paragraph(
                    f"Step {self.step:03d} at {ts} in '{title}': Click recorded at ({x},{y})")
                self.doc.add_paragraph("")
            except Exception as e2:
                print(f"Failed to record even simplified click: {e2}")

    # ── screenshot helper ──────────────────────────────────────────────
    def _capture_and_highlight(self, x: int, y: int, key=None) -> str:
        img = ImageGrab.grab()
        draw = ImageDraw.Draw(img)
        r = 30
        draw.ellipse((x - r, y - r, x + r, y + r), outline="red", width=5)
        if key is not None:
            kname = self._key_to_name(key)
            try:
                font = ImageFont.truetype("arial.ttf", 24)
            except Exception:
                font = ImageFont.load_default()
            bbox = draw.textbbox((0, 0), kname, font=font)
            tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
            box = [(x - r, y - r - th - 10), (x - r + tw + 10, y - r)]
            draw.rectangle(box, fill=(255, 255, 0, 128))
            draw.text((box[0][0] + 5, box[0][1] + 5), kname, fill="black", font=font)

        path = os.path.abspath(f"step_{self.step:04d}.png")
        img.save(path)
        self.outputs.append(path)
        return path


if __name__ == "__main__":
    if CTK_AVAILABLE:
        root = ctk.CTk()
    else:
        root = tk.Tk()
    
    StepRecorderApp(root)
    root.mainloop()
