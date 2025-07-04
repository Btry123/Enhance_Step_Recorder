import os
import json
import base64
from datetime import datetime
from typing import Dict, List, Any, Optional

from docx import Document
from docx.shared import Inches

# Optional imports
try:
    from docx2pdf import convert
    DOCX2PDF = True
except ImportError:
    DOCX2PDF = False


class BaseExporter:
    """Base class for all exporters."""
    
    def __init__(self, session_data):
        self.session = session_data
    
    def export(self, output_path: str) -> bool:
        """Export the recording to the specified path."""
        raise NotImplementedError("Subclasses must implement export()")


class DocxExporter(BaseExporter):
    """Export recording to Word document."""
    
    def export(self, output_path: str) -> bool:
        """Export to Word document."""
        try:
            doc = Document()
            doc.add_heading("Enhanced Step Recording", level=1)
            
            # Add metadata
            if self.session.start_time:
                doc.add_paragraph(f"Recording started at: {self.session.start_time.strftime('%Y-%m-%d %H:%M:%S')}")
            if self.session.end_time:
                doc.add_paragraph(f"Recording ended at: {self.session.end_time.strftime('%Y-%m-%d %H:%M:%S')}")
            
            duration = self.session.statistics.get('recording_duration', 0)
            doc.add_paragraph(f"Total duration: {duration:.1f} seconds")
            doc.add_paragraph(f"Total steps: {self.session.statistics.get('total_steps', 0)}")
            doc.add_paragraph("")
            
            # Add steps
            for step in self.session.steps:
                step_num = step.get('step_number', 0)
                timestamp = step.get('timestamp', datetime.now())
                step_type = step.get('type', 'unknown')
                content = step.get('content', '')
                description = step.get('description', '')  # AI description
                application = step.get('application', 'Unknown App')
                screenshot = step.get('screenshot')
                
                # Format timestamp
                ts_str = timestamp.strftime("%H:%M:%S")
                
                # Create step description with AI enhancement
                if step_type == 'note':
                    step_text = f"Step {step_num:03d} at {ts_str} in '{application}': NOTE - {content}"
                else:
                    # Use AI description if available, otherwise use original content
                    display_text = description if description else content
                    step_text = f"Step {step_num:03d} at {ts_str} in '{application}': {display_text}"
                
                doc.add_paragraph(step_text)
                
                # Add screenshot if available
                if screenshot and os.path.exists(screenshot):
                    try:
                        doc.add_picture(screenshot, width=Inches(5))
                    except Exception as e:
                        doc.add_paragraph(f"[Screenshot error: {e}]")
                
                doc.add_paragraph("")
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            return False


class HtmlExporter(BaseExporter):
    """Export recording to interactive HTML report."""
    
    def export(self, output_path: str) -> bool:
        """Export to HTML report."""
        try:
            html_content = self._generate_html()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
            
        except Exception as e:
            print(f"Error exporting to HTML: {e}")
            return False
    
    def _generate_html(self) -> str:
        """Generate HTML content."""
        # Calculate statistics
        stats = self.session.statistics
        duration = stats.get('recording_duration', 0)
        total_steps = stats.get('total_steps', 0)
        clicks = stats.get('clicks', 0)
        keystrokes = stats.get('keystrokes', 0)
        notes = stats.get('notes', 0)
        applications = list(stats.get('applications', set()))
        
        # Generate HTML
        html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Step Recording Report</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            margin: 0;
            padding: 20px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            box-shadow: 0 20px 40px rgba(0,0,0,0.1);
            overflow: hidden;
        }}
        .header {{
            background: linear-gradient(135deg, #4CAF50, #45a049);
            color: white;
            padding: 30px;
            text-align: center;
        }}
        .header h1 {{
            margin: 0;
            font-size: 2.5em;
            font-weight: 300;
        }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            padding: 30px;
            background: #f8f9fa;
        }}
        .stat-card {{
            background: white;
            padding: 20px;
            border-radius: 10px;
            text-align: center;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }}
        .stat-number {{
            font-size: 2em;
            font-weight: bold;
            color: #4CAF50;
        }}
        .stat-label {{
            color: #666;
            margin-top: 5px;
        }}
        .steps-container {{
            padding: 30px;
        }}
        .step {{
            background: white;
            margin-bottom: 20px;
            border-radius: 10px;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
            overflow: hidden;
        }}
        .step-header {{
            background: #f8f9fa;
            padding: 15px 20px;
            border-bottom: 1px solid #eee;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .step-number {{
            font-weight: bold;
            color: #4CAF50;
            font-size: 1.2em;
        }}
        .step-time {{
            color: #666;
            font-size: 0.9em;
        }}
        .step-content {{
            padding: 20px;
        }}
        .step-description {{
            margin-bottom: 15px;
            line-height: 1.6;
        }}
        .step-screenshot {{
            max-width: 100%;
            border-radius: 8px;
            box-shadow: 0 3px 10px rgba(0,0,0,0.2);
        }}
        .step-type {{
            display: inline-block;
            padding: 4px 8px;
            border-radius: 4px;
            font-size: 0.8em;
            font-weight: bold;
            text-transform: uppercase;
        }}
        .type-click {{
            background: #e3f2fd;
            color: #1976d2;
        }}
        .type-keystroke {{
            background: #f3e5f5;
            color: #7b1fa2;
        }}
        .type-note {{
            background: #fff3e0;
            color: #f57c00;
        }}
        .controls {{
            position: fixed;
            top: 20px;
            right: 20px;
            background: white;
            padding: 15px;
            border-radius: 10px;
            box-shadow: 0 5px 15px rgba(0,0,0,0.2);
        }}
        .btn {{
            background: #4CAF50;
            color: white;
            border: none;
            padding: 8px 15px;
            border-radius: 5px;
            cursor: pointer;
            margin: 2px;
            font-size: 0.9em;
        }}
        .btn:hover {{
            background: #45a049;
        }}
        @media (max-width: 768px) {{
            .stats-grid {{
                grid-template-columns: 1fr;
            }}
            .controls {{
                position: static;
                margin-bottom: 20px;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🎬 Step Recording Report</h1>
            <p>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
        
        <div class="stats-grid">
            <div class="stat-card">
                <div class="stat-number">{total_steps}</div>
                <div class="stat-label">Total Steps</div>
            </div>
            <div class="stat-card">
                <div class="stat-number">{duration:.1f}s</div>
                <div class="stat-label">Duration</div>
            </div>
            <div class="stat-card">
                <div class="stat-number">{clicks}</div>
                <div class="stat-label">Mouse Clicks</div>
            </div>
            <div class="stat-card">
                <div class="stat-number">{keystrokes}</div>
                <div class="stat-label">Keystrokes</div>
            </div>
            <div class="stat-card">
                <div class="stat-number">{notes}</div>
                <div class="stat-label">Notes</div>
            </div>
            <div class="stat-card">
                <div class="stat-number">{len(applications)}</div>
                <div class="stat-label">Applications</div>
            </div>
        </div>
        
        <div class="controls">
            <button class="btn" onclick="expandAll()">Expand All</button>
            <button class="btn" onclick="collapseAll()">Collapse All</button>
            <button class="btn" onclick="window.print()">Print</button>
        </div>
        
        <div class="steps-container">
            <h2>📋 Recording Steps</h2>
        """
        
        # Add steps
        for step in self.session.steps:
            step_num = step.get('step_number', 0)
            timestamp = step.get('timestamp', datetime.now())
            step_type = step.get('type', 'unknown')
            content = step.get('content', '')
            description = step.get('description', '')  # AI description
            application = step.get('application', 'Unknown App')
            screenshot = step.get('screenshot')
            
            ts_str = timestamp.strftime("%H:%M:%S")
            
            # Type badge
            type_class = f"type-{step_type}"
            type_label = step_type.upper()
            
            # Use AI description if available, otherwise use original content
            display_text = description if description else content
            ai_indicator = "🤖 " if description else ""
            
            html += f"""
            <div class="step">
                <div class="step-header">
                    <div>
                        <span class="step-number">Step {step_num:03d}</span>
                        <span class="step-type {type_class}">{type_label}</span>
                        {ai_indicator}
                    </div>
                    <div class="step-time">{ts_str} - {application}</div>
                </div>
                <div class="step-content">
                    <div class="step-description">{display_text}</div>
            """
            
            # Add screenshot if available
            if screenshot and os.path.exists(screenshot):
                try:
                    # Convert image to base64 for embedding
                    with open(screenshot, 'rb') as img_file:
                        img_data = base64.b64encode(img_file.read()).decode()
                    img_ext = os.path.splitext(screenshot)[1].lower()
                    mime_type = 'image/png' if img_ext == '.png' else 'image/jpeg'
                    
                    html += f'<img src="data:{mime_type};base64,{img_data}" class="step-screenshot" alt="Step {step_num} screenshot">'
                except Exception as e:
                    html += f'<p style="color: red;">[Screenshot error: {e}]</p>'
            
            html += """
                </div>
            </div>
            """
        
        # Close HTML
        html += """
        </div>
    </div>
    
    <script>
        function expandAll() {
            // Implementation for expand all
            console.log('Expand all clicked');
        }
        
        function collapseAll() {
            // Implementation for collapse all
            console.log('Collapse all clicked');
        }
    </script>
</body>
</html>
        """
        
        return html


class MarkdownExporter(BaseExporter):
    """Export recording to Markdown format."""
    
    def export(self, output_path: str) -> bool:
        """Export to Markdown."""
        try:
            markdown_content = self._generate_markdown()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return True
            
        except Exception as e:
            print(f"Error exporting to Markdown: {e}")
            return False
    
    def _generate_markdown(self) -> str:
        """Generate Markdown content."""
        stats = self.session.statistics
        duration = stats.get('recording_duration', 0)
        total_steps = stats.get('total_steps', 0)
        clicks = stats.get('clicks', 0)
        keystrokes = stats.get('keystrokes', 0)
        notes = stats.get('notes', 0)
        applications = list(stats.get('applications', set()))
        
        markdown = f"""# 🎬 Step Recording Report

**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 📊 Recording Statistics

| Metric | Value |
|--------|-------|
| Total Steps | {total_steps} |
| Duration | {duration:.1f} seconds |
| Mouse Clicks | {clicks} |
| Keystrokes | {keystrokes} |
| Notes | {notes} |
| Applications Used | {len(applications)} |

## 🖥️ Applications Used

{chr(10).join([f"- {app}" for app in applications])}

## 📋 Recording Steps

"""
        
        # Add steps
        for step in self.session.steps:
            step_num = step.get('step_number', 0)
            timestamp = step.get('timestamp', datetime.now())
            step_type = step.get('type', 'unknown')
            content = step.get('content', '')
            description = step.get('description', '')  # AI description
            application = step.get('application', 'Unknown App')
            screenshot = step.get('screenshot')
            
            ts_str = timestamp.strftime("%H:%M:%S")
            
            # Use AI description if available, otherwise use original content
            display_text = description if description else content
            ai_indicator = "🤖 " if description else ""
            
            # Step header
            markdown += f"### {ai_indicator}Step {step_num:03d} - {ts_str}\n\n"
            markdown += f"**Type:** {step_type.upper()}  \n"
            markdown += f"**Application:** {application}  \n"
            markdown += f"**Action:** {display_text}\n\n"
            
            # Screenshot reference
            if screenshot and os.path.exists(screenshot):
                screenshot_name = os.path.basename(screenshot)
                markdown += f"![Step {step_num}]({screenshot_name})\n\n"
            
            markdown += "---\n\n"
        
        return markdown


class JsonExporter(BaseExporter):
    """Export recording to JSON format."""
    
    def export(self, output_path: str) -> bool:
        """Export to JSON."""
        try:
            json_data = self._generate_json()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, indent=2, default=str)
            
            return True
            
        except Exception as e:
            print(f"Error exporting to JSON: {e}")
            return False
    
    def _generate_json(self) -> Dict[str, Any]:
        """Generate JSON data."""
        # Convert session to JSON-serializable format
        json_data = {
            'metadata': {
                'version': '2.0',
                'created_by': 'Enhanced Step Recorder',
                'exported_at': datetime.now().isoformat(),
                'start_time': self.session.start_time.isoformat() if self.session.start_time else None,
                'end_time': self.session.end_time.isoformat() if self.session.end_time else None
            },
            'statistics': self.session.statistics.copy(),
            'steps': []
        }
        
        # Convert applications set to list
        json_data['statistics']['applications'] = list(json_data['statistics']['applications'])
        
        # Add steps
        for step in self.session.steps:
            step_data = step.copy()
            step_data['timestamp'] = step_data['timestamp'].isoformat()
            json_data['steps'].append(step_data)
        
        return json_data


class ExportManager:
    """Manager for handling multiple export formats."""
    
    def __init__(self, session_data):
        self.session = session_data
        self.exporters = {
            'docx': DocxExporter(session_data),
            'html': HtmlExporter(session_data),
            'md': MarkdownExporter(session_data),
            'json': JsonExporter(session_data)
        }
    
    def export_all(self, base_path: str) -> Dict[str, bool]:
        """Export to all available formats."""
        results = {}
        
        for format_name, exporter in self.exporters.items():
            if format_name == 'docx':
                output_path = f"{base_path}.docx"
            else:
                output_path = f"{base_path}.{format_name}"
            
            try:
                success = exporter.export(output_path)
                results[format_name] = success
            except Exception as e:
                print(f"Error exporting to {format_name}: {e}")
                results[format_name] = False
        
        return results
    
    def export_to_format(self, format_name: str, output_path: str) -> bool:
        """Export to specific format."""
        if format_name not in self.exporters:
            print(f"Unsupported format: {format_name}")
            return False
        
        try:
            return self.exporters[format_name].export(output_path)
        except Exception as e:
            print(f"Error exporting to {format_name}: {e}")
            return False 